"""Build the ATS-friendly Word version of the CV.

ATS rules followed here: single column, no tables, no text boxes, no headers/footers,
no images, standard system font, literal bullet characters, standard section headings.
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import cv_content as C

FONT = "Calibri"
BODY = 10.5
DARK = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT = RGBColor(0x00, 0x33, 0x66)

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(BODY)
normal.font.color.rgb = DARK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0


def para(space_before=0, space_after=0, align=None, indent=None, hanging=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if indent is not None:
        pf.left_indent = Inches(indent)
    if hanging is not None:
        pf.first_line_indent = Inches(-hanging)
    return p


def run(p, text, bold=False, italic=False, size=BODY, color=None, caps=False):
    r = p.add_run(text.upper() if caps else text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = FONT
    r.font.color.rgb = color or DARK
    return r


def bottom_border(p, color="003366", size=8):
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    pPr.append(borders)


def heading(text):
    p = para(space_before=7.5, space_after=3.5)
    run(p, text, bold=True, size=11, color=ACCENT, caps=True)
    bottom_border(p)


def bullet(text, size=BODY):
    p = para(indent=0.18, hanging=0.18, space_after=1.5)
    run(p, "\u2022\t" + text, size=size)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(0.18))
    return p


# ---------------------------------------------------------------- header
p = para(align=WD_ALIGN_PARAGRAPH.CENTER)
run(p, C.NAME, bold=True, size=20, color=ACCENT)

p = para(align=WD_ALIGN_PARAGRAPH.CENTER, space_before=1)
run(p, C.HEADLINE, bold=True, size=10)

p = para(align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2)
run(p, C.CONTACT, size=9.5)

# ---------------------------------------------------------------- summary
heading("Professional Summary")
p = para()
run(p, C.SUMMARY)

# ---------------------------------------------------------------- competencies
heading("Core Competencies")
p = para()
run(p, "  \u2022  ".join(C.COMPETENCIES))

# ---------------------------------------------------------------- skills
heading("Technical Skills")
for label, detail in C.SKILLS:
    p = para(space_after=2)
    run(p, label + ": ", bold=True)
    run(p, detail)

# ---------------------------------------------------------------- experience
heading("Professional Experience")
for i, job in enumerate(C.EXPERIENCE):
    p = para(space_before=0 if i == 0 else 7)
    run(p, job["company"], bold=True)
    run(p, "  \u2013  " + job["location"])

    p = para(space_after=2)
    run(p, job["title"], bold=True, italic=True)
    run(p, "  |  " + job["dates"], italic=True)
    if job["note"]:
        run(p, "  (" + job["note"] + ")", italic=True, size=9.5)

    for b in job["bullets"]:
        bullet(b)

# ---------------------------------------------------------------- certifications
heading("Certifications")
for c in C.CERTIFICATIONS:
    bullet(c)

# ---------------------------------------------------------------- education
heading("Education")
for e in C.EDUCATION:
    bullet(e)

# ---------------------------------------------------------------- languages
heading("Languages")
p = para()
run(p, C.LANGUAGES)

doc.core_properties.title = f"{C.NAME} - CV"
doc.core_properties.author = C.NAME

OUT = "/Users/afseer/CV/Afseer KP - IT Team Lead CV.docx"
doc.save(OUT)
print("wrote", OUT)
